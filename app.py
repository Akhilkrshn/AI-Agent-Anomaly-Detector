import io
import os
import sys

from pathlib import Path
import numpy as np
import pandas as pd
import streamlit as st
import hashlib
from docx import Document
from dotenv import load_dotenv
from sklearn.ensemble import IsolationForest

load_dotenv()
sys.path.append(str(Path(__file__).parent.resolve()))

from ai_summary import generate_ai_executive_report, generate_executive_metrics
from txn_anomaly_detector import auto_clean_data

RULE_DEFINITIONS = {
    "dedup": "Duplicate Record Removal",
    "amount_impute": "Missing Amount Imputation (median fill)",
    "risk_correct": "Out-of-Range Risk Score Correction (median fill)",
}

st.set_page_config(
    page_title="Enterprise Anomaly Detector & Quality Hub",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

if "nav_page" not in st.session_state:
    st.session_state.nav_page = "1. Main Workstation"

if "ai_summary_text" not in st.session_state:
    st.session_state.ai_summary_text = ""

if "ai_summary_success" not in st.session_state:
    st.session_state.ai_summary_success = False

if "audit_logs" not in st.session_state:
    st.session_state.audit_logs = []

if "excel_export_bytes" not in st.session_state:
    st.session_state.excel_export_bytes = None

st.markdown(
    """
    <style>
    .stApp { background-color: #FFFFFF; color: #0F172A; }
    [data-testid="stMetricValue"] { font-size: 26px !important; font-weight: 700 !important; color: #1E293B !important; }
    [data-testid="stMetricLabel"] { font-size: 14px !important; color: #475569 !important; font-weight: 600 !important; }
    div[data-testid="metric-container"] { background-color: #F8FAFC; border: 1px solid #E2E8F0; border-radius: 10px; padding: 16px; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
    .audit-card { background-color: #F0F9FF; border-left: 5px solid #0284C7; padding: 12px 16px; border-radius: 6px; margin-bottom: 8px; color: #0369A1; font-size: 14px; font-weight: 600; }
    .export-card { background-color: #F8FAFC; border: 1px solid #E2E8F0; border-radius: 10px; padding: 20px; margin-bottom: 15px; }

    div[data-testid="stSidebar"] div.stButton > button {
        width: 100% !important; text-align: left !important; padding: 12px 16px !important;
        border-radius: 8px !important; border: 1px solid #CBD5E1 !important;
        background-color: #F8FAFC !important; color: #1E293B !important; font-weight: 600 !important;
        margin-bottom: 8px !important; display: block !important;
    }
    div[data-testid="stSidebar"] div.stButton > button:hover {
        background-color: #0284C7 !important; color: #FFFFFF !important; border-color: #0284C7 !important;
    }
    </style>
""",
    unsafe_allow_html=True,
)

@st.cache_data(show_spinner=False)
def load_uploaded_dataset(file_bytes_val: bytes, file_name: str) -> pd.DataFrame:
    ext = file_name.split(".")[-1].lower()
    file_bytes = io.BytesIO(file_bytes_val)

    if ext == "parquet":
        return pd.read_parquet(file_bytes)
    elif ext in ["xlsx", "xls"]:
        return pd.read_excel(file_bytes)
    elif ext == "csv":
        return pd.read_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

@st.cache_data(show_spinner=False)
def process_pipeline(_raw_df: pd.DataFrame, contamination: float, skip_rules: frozenset, dataset_fingerprint: str):
    clean_df, logs = auto_clean_data(_raw_df, skip_rules=skip_rules)
    features = [col for col in ["Amount", "Risk_Score"] if col in clean_df.columns]

    if features and not clean_df[features].isnull().any().any():
        model = IsolationForest(
            contamination=contamination,
            random_state=42,
            n_jobs=-1,
            max_samples=0.25,
        )
        clean_df["anomaly_score"] = model.fit_predict(clean_df[features])
        clean_df["Is_Anomaly"] = clean_df["anomaly_score"] == -1
        clean_df.drop(columns=["anomaly_score"], inplace=True)
    else:
        clean_df["Is_Anomaly"] = False
        if features and clean_df[features].isnull().any().any():
            logs.append({
                "id": "nan_warning",
                "message": "Anomaly detection skipped: missing values remain in Amount/Risk_Score after reverting a cleaning rule."
            })

    return clean_df, logs

def build_excel_file(dataframe: pd.DataFrame, max_export_rows: int = None) -> bytes:
    output = io.BytesIO()
    export_target = dataframe.head(max_export_rows) if max_export_rows else dataframe
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        export_target.to_excel(writer, sheet_name="Cleaned_Data", index=False)
        if "Is_Anomaly" in export_target.columns:
            anomalies_only = export_target[export_target["Is_Anomaly"]]
            if not anomalies_only.empty:
                anomalies_only.to_excel(writer, sheet_name="Flagged_Anomalies", index=False)
    return output.getvalue()

def escape_markdown_math(text: str) -> str:
    """
    Streamlit's st.markdown() auto-renders LaTeX for any text wrapped in single ($)
    or double ($$) dollar signs. Our AI-generated report text contains literal
    currency values (e.g. "$734,443.28"), which Streamlit was misinterpreting as
    LaTeX math delimiters -- swallowing the $ sign and rendering the number range
    that followed in a monospace math font with its own spacing rules. This escapes
    every literal dollar sign so Streamlit renders it as plain text instead.
    """
    return text.replace("$", "\\$")

def add_markdown_to_docx(doc: Document, markdown_text: str):
    lines = markdown_text.split("\n")
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue

        if line_str.startswith("#"):
            level = min(line_str.count("#"), 3)
            heading_text = line_str.lstrip("#").strip()
            doc.add_heading(heading_text, level=level)
            continue

        if line_str.startswith("- ") or line_str.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            line_str = line_str[2:]
        else:
            p = doc.add_paragraph()

        p.add_run(line_str)

def create_docx_report(metrics: dict, audit_logs: list, ai_insights: str = "") -> bytes:
    doc = Document()
    doc.add_heading("Executive Anomaly Audit Report", level=0)

    doc.add_heading("1. Executive Metrics Summary", level=1)
    doc.add_paragraph(f"Total Records Processed: {metrics['total_records']:,}")
    doc.add_paragraph(f"Total Anomalies Flagged: {metrics['total_anomalies']:,}")
    doc.add_paragraph(f"Anomaly Rate: {metrics['anomaly_rate_pct']}%")
    doc.add_paragraph(f"Total Anomalous Exposure: ${metrics['anomalous_exposure']:,.2f}")

    if ai_insights:
        doc.add_heading("2. AI Intelligence Insights (Groq)", level=1)
        add_markdown_to_docx(doc, ai_insights)

    doc.add_heading("3. Data Remediation Audit Log", level=1)
    if audit_logs:
        for log in audit_logs:
            if isinstance(log, dict):
                msg = log.get("message", str(log))
            else:
                msg = str(log)
            doc.add_paragraph(f"• {msg}")
    else:
        doc.add_paragraph("No remediation operations were required.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

#Sidebar Controls
st.sidebar.title("🛡️ Control Panel")
st.sidebar.markdown("---")
st.sidebar.markdown("### Navigation")

if st.sidebar.button("📊 1. Main Workstation", use_container_width=True):
    st.session_state.nav_page = "1. Main Workstation"
if st.sidebar.button("📈 2. Visual Analytics", use_container_width=True):
    st.session_state.nav_page = "2. Visual Analytics"
if st.sidebar.button("📥 3. Export Options", use_container_width=True):
    st.session_state.nav_page = "3. Export Options"
if st.sidebar.button("🤖 4. AI Executive Summary", use_container_width=True):
    st.session_state.nav_page = "4. AI Executive Summary"

st.sidebar.markdown("---")
st.sidebar.subheader("📂 Data Ingestion")
uploaded_file = st.sidebar.file_uploader(
    "Upload Dataset", type=["parquet", "xlsx", "xls", "csv"]
)
contamination_rate = st.sidebar.slider(
    "Model Contamination Rate", 0.001, 0.05, 0.01, step=0.001
)

if uploaded_file is None:
    st.title("Enterprise Anomaly Detector & Quality Hub")
    st.info("👈 Please upload a dataset (.parquet, .xlsx, .csv) in the Sidebar Control Panel to begin.")
    st.stop()

#Reset per-rule revert checkboxes whenever a new file is loaded
if st.session_state.get("last_uploaded") != uploaded_file.name:
    for rule_id in RULE_DEFINITIONS:
        st.session_state[f"revert_{rule_id}"] = False

active_skip_rules = frozenset(
    rule_id for rule_id in RULE_DEFINITIONS
    if st.session_state.get(f"revert_{rule_id}", False)
)

file_bytes = uploaded_file.getvalue()
dataset_fingerprint = hashlib.md5(file_bytes).hexdigest()

raw_df = load_uploaded_dataset(file_bytes, uploaded_file.name)

with st.spinner("Processing Dataset & Fitting Anomaly Engines..."):
    clean_df, logs = process_pipeline(raw_df, contamination_rate, active_skip_rules, dataset_fingerprint)

if (
    "df_processed" not in st.session_state
    or st.session_state.get("last_uploaded") != uploaded_file.name
    or st.session_state.get("last_contamination") != contamination_rate
    or st.session_state.get("last_skip_rules") != active_skip_rules
):
    st.session_state.df_processed = clean_df.copy()
    st.session_state.audit_logs = logs
    st.session_state.ai_summary_text = ""
    st.session_state.last_uploaded = uploaded_file.name
    st.session_state.last_contamination = contamination_rate
    st.session_state.last_skip_rules = active_skip_rules

df = st.session_state.df_processed
metrics = generate_executive_metrics(df)

current_page = str(st.session_state.nav_page).lower()

#1. MAIN WORKSTATION
if "1. main workstation" in current_page:
    st.title("Enterprise Anomaly Workstation")
    st.markdown("Automated detection, data remediation, and interactive human-in-the-loop review.")
    st.markdown("---")

    m_col1, m_col2, m_col3, m_col4 = st.columns(4)
    m_col1.metric("Total Transactions", f"{metrics['total_records']:,}")
    m_col2.metric("Flagged Anomalies", f"{metrics['total_anomalies']:,}")
    m_col3.metric("Anomaly Rate", f"{metrics['anomaly_rate_pct']}%")
    m_col4.metric("Anomalous Exposure", f"${metrics['anomalous_exposure']:,.2f}")

    st.markdown("---")
    st.subheader("🛠️ Data Quality & Cleaning Audit Log")

    log_left_col, log_right_col = st.columns([3, 1])

    with log_right_col:
        with st.expander("Auditor Intervention", expanded=True):
            applied_rule_ids = {
                log.get("id") for log in st.session_state.audit_logs
                if isinstance(log, dict) and log.get("id") in RULE_DEFINITIONS
            }
            if applied_rule_ids:
                for rule_id in RULE_DEFINITIONS:
                    if rule_id in applied_rule_ids:
                        st.checkbox(
                            f"Revert: {RULE_DEFINITIONS[rule_id]}",
                            key=f"revert_{rule_id}",
                        )
            else:
                st.caption("No reversible cleaning rules were triggered on this dataset.")

    with log_left_col:
        if st.session_state.audit_logs:
            for log_item in st.session_state.audit_logs:
                if isinstance(log_item, dict):
                    log_id = log_item.get("id")
                    msg = log_item.get("message", str(log_item))
                    is_reverted = st.session_state.get(f"revert_{log_id}", False) if log_id else False
                    status_html = " <span style='color:red;'>[OVERRIDDEN]</span>" if is_reverted else ""
                    st.markdown(
                        f"<div class='audit-card'>• {msg}{status_html}</div>",
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f"<div class='audit-card'>• {str(log_item)}</div>",
                        unsafe_allow_html=True,
                    )
        else:
            st.success("No data quality issues found.")

    st.markdown("---")
    st.subheader("🔍 Interactive Auditor Workstation")
    tab1, tab2 = st.tabs(["⚠️ Flagged Anomalies", "✅ Unflagged Data"])

    with tab1:
        anomalies_df = df[df["Is_Anomaly"]].copy()
        edited_anomalies = st.data_editor(
            anomalies_df,
            key="anomalies_editor",
            use_container_width=True,
            num_rows="dynamic",
            column_config={
                "Is_Anomaly": st.column_config.CheckboxColumn("Marked Anomaly?")
            },
        )
        if not edited_anomalies.equals(anomalies_df):
            remaining_df = df.drop(index=anomalies_df.index, errors="ignore")
            df = pd.concat([remaining_df, edited_anomalies], ignore_index=True)
            st.session_state.df_processed = df
            st.rerun()

    with tab2:
        clean_df_view = df[~df["Is_Anomaly"]].copy()

        if clean_df_view.empty:
            st.info("No unflagged records in the current dataset.")
        else:
            filter_col1, filter_col2, filter_col3 = st.columns(3)
            with filter_col1:
                channel_options = sorted(clean_df_view["Channel"].dropna().unique()) if "Channel" in clean_df_view.columns else []
                selected_channels = st.multiselect("Filter: Channel", channel_options, key="clean_filter_channel")
            with filter_col2:
                branch_options = sorted(clean_df_view["Source_Branch"].dropna().unique()) if "Source_Branch" in clean_df_view.columns else []
                selected_branches = st.multiselect("Filter: Source Branch", branch_options, key="clean_filter_branch")
            with filter_col3:
                if "Amount" in clean_df_view.columns:
                    amt_min, amt_max = float(clean_df_view["Amount"].min()), float(clean_df_view["Amount"].max())
                    selected_amount_range = st.slider(
                        "Filter: Amount Range", amt_min, amt_max, (amt_min, amt_max), key="clean_filter_amount"
                    )
                else:
                    selected_amount_range = None

            filtered_view = clean_df_view
            if selected_channels:
                filtered_view = filtered_view[filtered_view["Channel"].isin(selected_channels)]
            if selected_branches:
                filtered_view = filtered_view[filtered_view["Source_Branch"].isin(selected_branches)]
            if selected_amount_range:
                filtered_view = filtered_view[
                    (filtered_view["Amount"] >= selected_amount_range[0])
                    & (filtered_view["Amount"] <= selected_amount_range[1])
                ]

            page_size = 1000
            total_rows = len(filtered_view)
            total_pages = max(1, (total_rows - 1) // page_size + 1)

            #Clamp any stale page number BEFORE the widget renders, or Streamlit
            #raises an error when a filter shrinks total_pages below the stored value.
            if st.session_state.get("clean_page_number", 1) > total_pages:
                st.session_state["clean_page_number"] = total_pages

            page_col, info_col = st.columns([1, 3])
            with page_col:
                page_number = st.number_input(
                    "Page", min_value=1, max_value=total_pages, step=1, key="clean_page_number"
                )
            with info_col:
                st.caption(f"Showing page {page_number} of {total_pages} — {total_rows:,} unflagged records match current filters.")

            start_idx = (page_number - 1) * page_size
            clean_page = filtered_view.iloc[start_idx : start_idx + page_size]

            edited_clean = st.data_editor(
                clean_page,
                key="clean_editor",
                use_container_width=True,
                num_rows="dynamic",
                column_config={
                    "Is_Anomaly": st.column_config.CheckboxColumn("Marked Anomaly?")
                },
            )
            if not edited_clean.equals(clean_page):
                remaining_df = df.drop(index=clean_page.index, errors="ignore")
                df = pd.concat([remaining_df, edited_clean], ignore_index=True)
                st.session_state.df_processed = df
                st.rerun()

#2. VISUAL ANALYTICS
elif "2. visual analytics" in current_page:
    st.title("📈 Visual Analytics")
    st.markdown("---")

    if df is not None and not df.empty:
        st.subheader("Risk Score Distribution")
        if "Risk_Score" in df.columns:
            # Human-readable bin labels instead of JSON interval strings
            counts, bin_edges = np.histogram(df["Risk_Score"].dropna(), bins=10)
            bin_labels = [f"{bin_edges[i]:.2f} - {bin_edges[i+1]:.2f}" for i in range(len(counts))]
            chart_df = pd.DataFrame({"Transactions": counts}, index=bin_labels)
            st.bar_chart(chart_df)
        else:
            st.info("Risk_Score column missing.")

        st.subheader("High Amount Anomalies")
        if "Is_Anomaly" in df.columns and "Amount" in df.columns:
            st.scatter_chart(df, x="Amount", y="Risk_Score", color="Is_Anomaly")
        else:
            st.info("Amount or Is_Anomaly columns missing.")
    else:
        st.info("👈 Please upload a dataset in the sidebar to load visual analytics.")

#3. EXPORT OPTIONS
elif "3. export options" in current_page:
    st.title("📥 Export Data Center")
    st.markdown("---")

    exp_c1, exp_c2 = st.columns(2)
    with exp_c1:
        st.markdown("<div class='export-card'>", unsafe_allow_html=True)
        st.markdown("#### 1. Export Parquet File")
        parquet_buf = io.BytesIO()
        df.to_parquet(parquet_buf, index=False)
        st.download_button(
            "📦 Download Parquet",
            parquet_buf.getvalue(),
            "anomaly_processed.parquet",
            "application/octet-stream",
            use_container_width=True,
        )
        st.caption(f"{len(df):,} rows ready.")
        st.markdown("</div>", unsafe_allow_html=True)

    with exp_c2:
        st.markdown("<div class='export-card'>", unsafe_allow_html=True)
        st.markdown("#### 2. Export Excel File")
        st.caption("Excel exports can take up to a minute on large datasets -- generate on demand below.")

        if st.button("⚙️ Generate Excel Export", use_container_width=True):
            with st.spinner(f"Building Excel workbook for {len(df):,} rows..."):
                st.session_state.excel_export_bytes = build_excel_file(df)
                st.session_state.excel_export_row_count = len(df)

        if st.session_state.get("excel_export_bytes"):
            st.download_button(
                "📊 Download Excel",
                st.session_state.excel_export_bytes,
                "cleaned_anomalies.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
            st.caption(f"{st.session_state.excel_export_row_count:,} rows in this file.")
        st.markdown("</div>", unsafe_allow_html=True)


#4. AI EXECUTIVE SUMMARY
elif "4. ai executive summary" in current_page:
    st.title("🤖 AI Executive Summary Generator")
    st.markdown("---")

    if st.button("⚡ Generate AI Summary with Groq", use_container_width=True):
        with st.spinner("Querying Groq API..."):
            success, result_text = generate_ai_executive_report(
                df, st.session_state.audit_logs
            )
            st.session_state.ai_summary_success = success
            st.session_state.ai_summary_text = result_text

    if st.session_state.ai_summary_text:
        st.markdown("---")
        if st.session_state.ai_summary_success:
            st.markdown(escape_markdown_math(st.session_state.ai_summary_text))

            docx_bytes = create_docx_report(
                metrics,
                st.session_state.audit_logs,
                st.session_state.ai_summary_text,
            )
            st.download_button(
                "📄 Download Word Report (.docx)",
                docx_bytes,
                "Executive_Anomaly_Report.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.error(f"AI summary generation failed: {st.session_state.ai_summary_text}")